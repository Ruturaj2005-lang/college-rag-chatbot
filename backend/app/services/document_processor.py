import os
import re
from typing import List, Dict, Any
from app.core.logging import logger

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import docx
except ImportError:
    docx = None


class ExtractedPage:
    def __init__(self, page_number: int, text: str, heading: str = ""):
        self.page_number = page_number
        self.text = text
        self.heading = heading


def clean_extracted_text(text: str) -> str:
    """Cleans up text while preserving paragraph breaks and structure."""
    if not text:
        return ""
    # Normalize carriage returns
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Clean each line's whitespace
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    text = "\n".join(lines)
    # Replace 3 or more newlines with 2 newlines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_from_pdf(file_path: str) -> List[ExtractedPage]:
    """Extracts text page by page from a PDF file using PyMuPDF."""
    if fitz is None:
        raise RuntimeError("PyMuPDF (fitz) is not installed.")
    
    pages: List[ExtractedPage] = []
    doc = fitz.open(file_path)
    try:
        for page_idx, page in enumerate(doc):
            text = page.get_text("text")
            cleaned = clean_extracted_text(text)
            if cleaned:
                pages.append(ExtractedPage(page_number=page_idx + 1, text=cleaned))
    finally:
        doc.close()
    
    if not pages:
        logger.warning(f"No extractable text found in PDF: {file_path}")
    return pages


def extract_from_docx(file_path: str) -> List[ExtractedPage]:
    """Extracts text and table content from a DOCX file using python-docx."""
    if docx is None:
        raise RuntimeError("python-docx is not installed.")
    
    doc = docx.Document(file_path)
    full_paragraphs = []
    
    # Process headings and paragraphs
    for p in doc.paragraphs:
        p_text = p.text.strip()
        if p_text:
            if p.style.name.startswith("Heading"):
                full_paragraphs.append(f"\n### {p_text}\n")
            else:
                full_paragraphs.append(p_text)
                
    # Process tables if any
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(row_data):
                table_rows.append(" | ".join(row_data))
        if table_rows:
            full_paragraphs.append("\n" + "\n".join(table_rows) + "\n")
            
    combined_text = "\n\n".join(full_paragraphs)
    cleaned = clean_extracted_text(combined_text)
    
    # Approximate page estimation (~3000 chars per page if not paginated)
    if len(cleaned) <= 3000:
        return [ExtractedPage(page_number=1, text=cleaned)]
    else:
        pages = []
        parts = [cleaned[i:i + 3000] for i in range(0, len(cleaned), 3000)]
        for idx, part in enumerate(parts):
            pages.append(ExtractedPage(page_number=idx + 1, text=part.strip()))
        return pages


def extract_from_txt(file_path: str) -> List[ExtractedPage]:
    """Extracts text from a plain TXT file with robust encoding detection."""
    content = ""
    for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            with open(file_path, "r", encoding=enc) as f:
                content = f.read()
            break
        except UnicodeDecodeError:
            continue
            
    cleaned = clean_extracted_text(content)
    if len(cleaned) <= 3000:
        return [ExtractedPage(page_number=1, text=cleaned)]
    else:
        pages = []
        parts = [cleaned[i:i + 3000] for i in range(0, len(cleaned), 3000)]
        for idx, part in enumerate(parts):
            pages.append(ExtractedPage(page_number=idx + 1, text=part.strip()))
        return pages


def process_document_file(file_path: str, file_type: str) -> List[ExtractedPage]:
    """Routes the file to the appropriate extractor based on extension/type."""
    file_type = file_type.lower().strip(".")
    if file_type == "pdf":
        return extract_from_pdf(file_path)
    elif file_type in ["docx", "doc"]:
        return extract_from_docx(file_path)
    elif file_type in ["txt", "text", "md"]:
        return extract_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}. Supported types: PDF, DOCX, TXT")
